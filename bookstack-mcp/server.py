import base64
import contextlib
import os
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
import re

import requests
import uvicorn
from bs4 import BeautifulSoup
from mcp.server.fastmcp import FastMCP
from mcp.server.transport_security import TransportSecuritySettings
from pypdf import PdfReader
from starlette.applications import Starlette
from starlette.responses import JSONResponse
from starlette.routing import Mount, Route


BOOKSTACK_BASE_URL = os.getenv("BOOKSTACK_BASE_URL", "http://bookstack").rstrip("/")
BOOKSTACK_TOKEN_ID = os.getenv("BOOKSTACK_TOKEN_ID")
BOOKSTACK_TOKEN_SECRET = os.getenv("BOOKSTACK_TOKEN_SECRET")

MCP_HOST = os.getenv("MCP_HOST", "0.0.0.0")
MCP_PORT = int(os.getenv("MCP_PORT", "8000"))

# External/link attachment fetching is OFF by default for security.
# Uploaded BookStack PDFs do not need this. They are read through /api/attachments/{id}.
FETCH_EXTERNAL_ATTACHMENTS = os.getenv("FETCH_EXTERNAL_ATTACHMENTS", "false").lower() in {
    "1",
    "true",
    "yes",
    "on",
}
EXTERNAL_ATTACHMENT_ALLOWED_HOSTS = {
    host.strip().lower()
    for host in os.getenv("EXTERNAL_ATTACHMENT_ALLOWED_HOSTS", "").split(",")
    if host.strip()
}
MAX_EXTERNAL_ATTACHMENT_BYTES = int(os.getenv("MAX_EXTERNAL_ATTACHMENT_BYTES", "26214400"))

# Extra Host/Origin allow-list entries for MCP DNS-rebinding protection.
# Example:
# MCP_ALLOWED_HOSTS=bookstack-mcp,bookstack-mcp:8000,mdhbookstack.duckdns.org,mdhbookstack.duckdns.org:*
# MCP_ALLOWED_ORIGINS=http://langflow:7860,https://your-langflow-domain.duckdns.org
EXTRA_MCP_ALLOWED_HOSTS = [
    host.strip()
    for host in os.getenv("MCP_ALLOWED_HOSTS", "").split(",")
    if host.strip()
]
EXTRA_MCP_ALLOWED_ORIGINS = [
    origin.strip()
    for origin in os.getenv("MCP_ALLOWED_ORIGINS", "").split(",")
    if origin.strip()
]


mcp = FastMCP(
    name="MDH BookStack MCP Server",
    stateless_http=False,
    json_response=False,
    transport_security=TransportSecuritySettings(
        enable_dns_rebinding_protection=True,
        allowed_hosts=[
            "localhost:*",
            "127.0.0.1:*",
            "bookstack-mcp",
            "bookstack-mcp:*",
            "bookstack-mcp:8000",
            *EXTRA_MCP_ALLOWED_HOSTS,
        ],
        allowed_origins=[
            "http://localhost:*",
            "http://127.0.0.1:*",
            "http://langflow:*",
            "http://langflow:7860",
            *EXTRA_MCP_ALLOWED_ORIGINS,
        ],
    ),
)


def get_bookstack_headers() -> dict[str, str]:
    if not BOOKSTACK_TOKEN_ID or not BOOKSTACK_TOKEN_SECRET:
        raise RuntimeError(
            "Missing BOOKSTACK_TOKEN_ID or BOOKSTACK_TOKEN_SECRET environment variable."
        )

    return {
        "Authorization": f"Token {BOOKSTACK_TOKEN_ID}:{BOOKSTACK_TOKEN_SECRET}",
        "Accept": "application/json",
    }


def call_bookstack_api(
    path: str,
    params: dict[str, Any] | None = None,
) -> dict[str, Any]:
    url = f"{BOOKSTACK_BASE_URL}{path}"

    response = requests.get(
        url,
        headers=get_bookstack_headers(),
        params=params,
        timeout=60,
    )

    if response.status_code >= 400:
        raise RuntimeError(
            {
                "error": "BookStack API request failed",
                "url": url,
                "status_code": response.status_code,
                "response": response.text,
            }
        )

    return response.json()


def html_to_clean_text(html: str) -> str:
    soup = BeautifulSoup(html or "", "html.parser")

    for tag in soup(["script", "style"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)
    lines = []

    for line in text.splitlines():
        clean_line = line.strip()
        if clean_line:
            lines.append(clean_line)

    return "\n".join(lines)


def simplify_page(page: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": page.get("id"),
        "name": page.get("name"),
        "slug": page.get("slug"),
        "book_id": page.get("book_id"),
        "chapter_id": page.get("chapter_id"),
        "url": page.get("url"),
        "created_at": page.get("created_at"),
        "updated_at": page.get("updated_at"),
    }


def get_uploaded_to_id(uploaded_to: Any) -> int | None:
    if isinstance(uploaded_to, dict):
        uploaded_to = uploaded_to.get("id")

    if uploaded_to is None:
        return None

    try:
        return int(uploaded_to)
    except (TypeError, ValueError):
        return None


def list_all_attachments_internal(max_attachments: int = 10000) -> list[dict[str, Any]]:
    max_attachments = max(1, min(max_attachments, 10000))
    collected_attachments: list[dict[str, Any]] = []

    offset = 0
    batch_size = 500

    while len(collected_attachments) < max_attachments:
        result = call_bookstack_api(
            "/api/attachments",
            {
                "count": batch_size,
                "offset": offset,
            },
        )

        attachments = result.get("data", [])

        if not attachments:
            break

        collected_attachments.extend(attachments)
        total = result.get("total")

        if isinstance(total, int) and offset + len(attachments) >= total:
            break

        if len(attachments) < batch_size:
            break

        offset += batch_size

    return collected_attachments[:max_attachments]


def list_attachments_for_page_internal(page_id: int) -> list[dict[str, Any]]:
    attachments = list_all_attachments_internal(max_attachments=10000)
    return [
        attachment
        for attachment in attachments
        if get_uploaded_to_id(attachment.get("uploaded_to")) == int(page_id)
    ]


def group_attachments_by_page(
    attachments: list[dict[str, Any]],
) -> dict[int, list[dict[str, Any]]]:
    grouped: dict[int, list[dict[str, Any]]] = {}

    for attachment in attachments:
        page_id = get_uploaded_to_id(attachment.get("uploaded_to"))

        if page_id is None:
            continue

        grouped.setdefault(page_id, []).append(attachment)

    return grouped


def extract_attachment_ids_from_text(value: str) -> list[int]:
    """
    Detect BookStack attachment IDs from page HTML/link text.

    Common BookStack file links look like /attachments/{id}.
    This fallback helps when the PDF is inserted as a link inside the page body
    instead of being returned by /api/attachments as uploaded_to=<page_id>.
    """
    if not value:
        return []

    patterns = [
        r"/api/attachments/(\d+)",
        r"/attachments/(\d+)",
        r"attachment[_-]?id[\s=:&quot;'\"]+(\d+)",
        r"data-attachment-id=[&quot;'\"]?(\d+)",
    ]

    found: set[int] = set()
    for pattern in patterns:
        for match in re.findall(pattern, value, flags=re.IGNORECASE):
            try:
                found.add(int(match))
            except (TypeError, ValueError):
                continue

    return sorted(found)


def extract_attachment_links_from_html(html: str) -> list[dict[str, Any]]:
    """Return links from page HTML that look like BookStack/PDF attachments."""
    soup = BeautifulSoup(html or "", "html.parser")
    links: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    for tag in soup.find_all(["a", "iframe", "embed", "object"]):
        href = (
            tag.get("href")
            or tag.get("src")
            or tag.get("data")
            or tag.get("data-url")
            or ""
        )
        label = tag.get_text(" ", strip=True) or tag.get("title") or tag.get("alt") or ""
        href_text = str(href).strip()
        label_text = str(label).strip()

        combined = f"{href_text} {label_text}"
        detected_ids = extract_attachment_ids_from_text(combined)
        looks_like_file = bool(
            detected_ids
            or ".pdf" in combined.lower()
            or ".docx" in combined.lower()
            or ".xlsx" in combined.lower()
            or ".txt" in combined.lower()
            or ".csv" in combined.lower()
        )

        if not looks_like_file:
            continue

        key = (href_text, label_text)
        if key in seen:
            continue

        seen.add(key)
        links.append(
            {
                "tag": tag.name,
                "href": href_text,
                "label": label_text,
                "detected_attachment_ids": detected_ids,
            }
        )

    # Some editors store attachment metadata outside normal links.
    html_detected_ids = extract_attachment_ids_from_text(html or "")
    if html_detected_ids and not any(link["detected_attachment_ids"] for link in links):
        links.append(
            {
                "tag": "html",
                "href": "",
                "label": "attachment ids found in raw html",
                "detected_attachment_ids": html_detected_ids,
            }
        )

    return links


def get_linked_attachment_stubs_from_page_html(
    page_id: int,
    html: str,
    existing_attachment_ids: set[int] | None = None,
) -> list[dict[str, Any]]:
    """Build attachment stubs from attachment IDs discovered in the page HTML."""
    existing_attachment_ids = existing_attachment_ids or set()
    stubs: list[dict[str, Any]] = []

    for link in extract_attachment_links_from_html(html):
        for attachment_id in link.get("detected_attachment_ids", []):
            if attachment_id in existing_attachment_ids:
                continue

            existing_attachment_ids.add(attachment_id)
            stubs.append(
                {
                    "id": attachment_id,
                    "name": link.get("label") or f"linked-attachment-{attachment_id}",
                    "uploaded_to": page_id,
                    "external": False,
                    "source": "page_html_link",
                    "links": {"html_href": link.get("href")},
                }
            )

    return stubs


def get_extension_from_name_or_url(name: str = "", url: str = "") -> str:
    candidates = []

    if name:
        candidates.append(name)

    if url:
        parsed = urlparse(url)
        candidates.append(parsed.path)

    for candidate in candidates:
        suffix = Path(candidate).suffix.lower().strip(".")
        if suffix:
            return suffix

    return ""


def extension_from_content_type(content_type: str | None) -> str:
    content_type = (content_type or "").split(";", 1)[0].lower().strip()

    mapping = {
        "application/pdf": "pdf",
        "text/plain": "txt",
        "text/markdown": "md",
        "text/csv": "csv",
        "application/json": "json",
        "application/xml": "xml",
        "text/xml": "xml",
        "text/html": "html",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    }

    return mapping.get(content_type, "")


def get_attachment_extension(attachment: dict[str, Any]) -> str:
    extension = (attachment.get("extension") or "").lower().strip(".")

    if extension:
        return extension

    name = attachment.get("name") or ""
    content = str(attachment.get("content") or "")
    return get_extension_from_name_or_url(name=name, url=content)


def decode_attachment_content(attachment: dict[str, Any]) -> bytes:
    """
    For uploaded files, BookStack returns base64 content.
    For external/link attachments, content is only the URL.
    """
    content = attachment.get("content")

    if not content:
        return b""

    if attachment.get("external"):
        return str(content).encode("utf-8", errors="ignore")

    content_text = str(content).strip()

    if content_text.startswith("data:") and "," in content_text:
        content_text = content_text.split(",", 1)[1]

    return base64.b64decode(content_text, validate=False)


def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract real text from a PDF file.
    This works for normal text PDFs. Scanned image PDFs need OCR.
    """
    if not file_bytes:
        return ""

    pdf_reader = PdfReader(BytesIO(file_bytes))
    text_parts = []

    for page_index, pdf_page in enumerate(pdf_reader.pages, start=1):
        page_text = pdf_page.extract_text() or ""

        if page_text.strip():
            text_parts.append(
                f"\n--- PDF Page {page_index} ---\n{page_text.strip()}"
            )

    return "\n".join(text_parts).strip()


def extract_text_file_content(file_bytes: bytes) -> str:
    if not file_bytes:
        return ""

    return file_bytes.decode("utf-8", errors="ignore").strip()


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "DOCX attachment found, but python-docx is not installed."

    document = Document(BytesIO(file_bytes))
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_xlsx_text(file_bytes: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "XLSX attachment found, but openpyxl is not installed."

    workbook = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []

    for sheet in workbook.worksheets:
        parts.append(f"--- XLSX Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                parts.append(" | ".join(values))

    workbook.close()
    return "\n".join(parts).strip()


def extract_bytes_by_extension(
    file_bytes: bytes,
    extension: str,
    attachment_name: str,
) -> str:
    extension = (extension or "").lower().strip(".")

    if not file_bytes:
        return ""

    if extension == "pdf":
        return extract_pdf_text(file_bytes)

    if extension in ["txt", "md", "csv", "json", "xml"]:
        return extract_text_file_content(file_bytes)

    if extension in ["html", "htm"]:
        return html_to_clean_text(extract_text_file_content(file_bytes))

    if extension == "docx":
        return extract_docx_text(file_bytes)

    if extension == "xlsx":
        return extract_xlsx_text(file_bytes)

    return (
        f"Attachment '{attachment_name}' was found, "
        f"but text extraction is not supported for extension '{extension}'."
    )


def download_external_attachment(url: str) -> tuple[bytes, str, str]:
    parsed = urlparse(url)
    host = (parsed.hostname or "").lower()

    if parsed.scheme not in {"http", "https"}:
        raise ValueError(f"External attachment URL scheme is not allowed: {parsed.scheme}")

    if not host:
        raise ValueError("External attachment URL has no host.")

    if EXTERNAL_ATTACHMENT_ALLOWED_HOSTS and host not in EXTERNAL_ATTACHMENT_ALLOWED_HOSTS:
        raise ValueError(
            f"External attachment host '{host}' is not allowed. "
            "Add it to EXTERNAL_ATTACHMENT_ALLOWED_HOSTS if you trust it."
        )

    response = requests.get(url, timeout=60, stream=True)
    response.raise_for_status()

    chunks = []
    total_size = 0

    for chunk in response.iter_content(chunk_size=1024 * 1024):
        if not chunk:
            continue

        total_size += len(chunk)

        if total_size > MAX_EXTERNAL_ATTACHMENT_BYTES:
            raise ValueError(
                f"External attachment is larger than MAX_EXTERNAL_ATTACHMENT_BYTES={MAX_EXTERNAL_ATTACHMENT_BYTES}."
            )

        chunks.append(chunk)

    content_type = response.headers.get("Content-Type", "")
    return b"".join(chunks), content_type, response.url


def extract_attachment_text(
    attachment_id: int,
    attachment_stub: dict[str, Any] | None = None,
) -> dict[str, Any]:
    """
    Read one BookStack attachment and extract readable text.

    If BookStack cannot read the stored file, skip this attachment instead of
    stopping get_all_pages_as_rag_text.
    """
    try:
        attachment = call_bookstack_api(f"/api/attachments/{attachment_id}")
    except Exception as exc:
        attachment_stub = attachment_stub or {}
        return {
            "id": attachment_id,
            "name": attachment_stub.get("name") or f"attachment-{attachment_id}",
            "extension": get_attachment_extension(attachment_stub),
            "external": attachment_stub.get("external"),
            "uploaded_to": attachment_stub.get("uploaded_to"),
            "created_at": attachment_stub.get("created_at"),
            "updated_at": attachment_stub.get("updated_at"),
            "links": attachment_stub.get("links"),
            "source_url": attachment_stub.get("content") if attachment_stub.get("external") else None,
            "text": "",
            "extraction_error": str(exc),
            "skipped": True,
        }

    attachment_name = attachment.get("name") or f"attachment-{attachment_id}"
    extension = get_attachment_extension(attachment)
    is_external = bool(attachment.get("external"))

    extracted_text = ""
    extraction_error = None
    skipped = False

    try:
        if is_external:
            external_url = str(attachment.get("content") or "").strip()

            if FETCH_EXTERNAL_ATTACHMENTS and external_url:
                external_bytes, content_type, final_url = download_external_attachment(external_url)
                external_extension = (
                    extension
                    or get_extension_from_name_or_url(name=attachment_name, url=final_url)
                    or extension_from_content_type(content_type)
                )
                extracted_text = extract_bytes_by_extension(
                    external_bytes,
                    external_extension,
                    attachment_name,
                )
            else:
                skipped = True
                extracted_text = ""
        else:
            file_bytes = decode_attachment_content(attachment)
            extracted_text = extract_bytes_by_extension(
                file_bytes,
                extension,
                attachment_name,
            )

        if not extracted_text.strip():
            skipped = True
            extracted_text = ""

    except Exception as exc:
        extraction_error = str(exc)
        skipped = True
        extracted_text = ""

    return {
        "id": attachment.get("id"),
        "name": attachment_name,
        "extension": extension,
        "external": is_external,
        "uploaded_to": attachment.get("uploaded_to"),
        "created_at": attachment.get("created_at"),
        "updated_at": attachment.get("updated_at"),
        "links": attachment.get("links"),
        "source_url": attachment.get("content") if is_external else None,
        "text": extracted_text,
        "extraction_error": extraction_error,
        "skipped": skipped,
    }


def build_page_text_with_attachments(
    page: dict[str, Any],
    page_attachments: list[dict[str, Any]] | None = None,
) -> tuple[str, list[dict[str, Any]]]:
    page_id = page.get("id")
    html = page.get("html") or ""
    page_text = html_to_clean_text(html)

    if page_id is None:
        return page_text, []

    if page_attachments is None:
        attachments = list_attachments_for_page_internal(int(page_id))
    else:
        attachments = list(page_attachments)

    # Fallback: If the PDF/file is inserted as a link in the page body,
    # it may not appear under /api/attachments with uploaded_to=<page_id>.
    # Detect /attachments/{id} links from the page HTML and read those IDs too.
    existing_attachment_ids = {
        int(attachment.get("id"))
        for attachment in attachments
        if attachment.get("id") is not None
    }
    attachments.extend(
        get_linked_attachment_stubs_from_page_html(
            page_id=int(page_id),
            html=html,
            existing_attachment_ids=existing_attachment_ids,
        )
    )

    extracted_attachments = []

    for attachment in attachments:
        attachment_id = attachment.get("id")

        if attachment_id is None:
            continue

        extracted_attachment = extract_attachment_text(
            int(attachment_id),
            attachment_stub=attachment,
        )
        extracted_attachments.append(extracted_attachment)

        if extracted_attachment.get("skipped"):
            continue

        attachment_text = extracted_attachment.get("text", "")

        if attachment_text.strip():
            page_text = (
                f"{page_text}\n\n"
                f"--- Attachment: {extracted_attachment.get('name')} "
                f"[id={extracted_attachment.get('id')}, "
                f"extension={extracted_attachment.get('extension')}, "
                f"external={extracted_attachment.get('external')}] ---\n"
                f"{attachment_text}"
            )

    return page_text.strip(), extracted_attachments


def get_page_internal(
    page_id: int,
    include_attachments: bool = True,
    page_attachments: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    page = call_bookstack_api(f"/api/pages/{page_id}")

    if include_attachments:
        text, attachments = build_page_text_with_attachments(
            page=page,
            page_attachments=page_attachments,
        )
    else:
        text = html_to_clean_text(page.get("html") or "")
        attachments = []

    attachment_text_count = sum(
        1
        for attachment in attachments
        if not attachment.get("skipped") and str(attachment.get("text") or "").strip()
    )
    attachment_skipped_count = sum(1 for attachment in attachments if attachment.get("skipped"))

    return {
        "id": page.get("id"),
        "name": page.get("name"),
        "slug": page.get("slug"),
        "book_id": page.get("book_id"),
        "chapter_id": page.get("chapter_id"),
        "url": page.get("url"),
        "created_at": page.get("created_at"),
        "updated_at": page.get("updated_at"),
        "text": text,
        "attachment_count": len(attachments),
        "attachment_text_count": attachment_text_count,
        "attachment_skipped_count": attachment_skipped_count,
        "attachments": attachments,
    }


def format_page_for_rag(page: dict[str, Any]) -> str:
    page_name = page.get("name") or "Untitled Page"
    page_id = page.get("id")
    page_url = page.get("url") or ""
    updated_at = page.get("updated_at") or ""
    page_text = page.get("text") or ""

    return "\n".join(
        [
            "==============================",
            f"PAGE TITLE: {page_name}",
            f"PAGE ID: {page_id}",
            f"PAGE URL: {page_url}",
            f"UPDATED AT: {updated_at}",
            "CONTENT:",
            page_text,
            "==============================",
        ]
    )


@mcp.tool()
def bookstack_status() -> dict[str, Any]:
    """
    Check whether the MCP server can connect to the BookStack API.
    """
    result = call_bookstack_api("/api/pages", {"count": 1})

    return {
        "status": "connected",
        "bookstack_base_url": BOOKSTACK_BASE_URL,
        "sample_result_keys": list(result.keys()),
    }


@mcp.tool()
def list_pages(count: int = 20, offset: int = 0) -> dict[str, Any]:
    """
    List BookStack pages.
    """
    count = max(1, min(count, 100))
    offset = max(0, offset)

    result = call_bookstack_api(
        "/api/pages",
        {
            "count": count,
            "offset": offset,
        },
    )

    pages = result.get("data", [])

    return {
        "total": result.get("total"),
        "count": count,
        "offset": offset,
        "pages": [simplify_page(page) for page in pages],
    }


@mcp.tool()
def get_page(page_id: int, include_attachments: bool = True) -> dict[str, Any]:
    """
    Get one BookStack page by page ID.
    If include_attachments is true, uploaded PDF text is added to the page text.
    """
    return get_page_internal(
        page_id=page_id,
        include_attachments=include_attachments,
        page_attachments=None,
    )


@mcp.tool()
def get_page_as_rag_text(
    page_id: int,
    include_attachments: bool = True,
) -> dict[str, Any]:
    """
    Get one BookStack page as clean text for Langflow RAG ingestion.
    """
    page = get_page_internal(
        page_id=page_id,
        include_attachments=include_attachments,
        page_attachments=None,
    )

    return {
        "page_id": page.get("id"),
        "page_name": page.get("name"),
        "page_url": page.get("url"),
        "attachment_count": page.get("attachment_count"),
        "attachment_text_count": page.get("attachment_text_count"),
        "attachment_skipped_count": page.get("attachment_skipped_count"),
        "rag_text": format_page_for_rag(page),
    }


@mcp.tool()
def search_bookstack(query: str, count: int = 10) -> dict[str, Any]:
    """
    Search BookStack content using the BookStack search API.
    Note: BookStack's own search may not index attachment text. Use RAG ingestion for that.
    """
    if not query.strip():
        raise ValueError("query cannot be empty")

    count = max(1, min(count, 50))

    result = call_bookstack_api(
        "/api/search",
        {
            "query": query,
            "count": count,
        },
    )

    return result


@mcp.tool()
def debug_page_attachment_detection(page_id: int) -> dict[str, Any]:
    """
    Debug why a page is not reading a PDF/file.

    Shows:
    - attachments found through BookStack /api/attachments
    - file-like links found inside the page HTML
    - attachment IDs detected from those HTML links
    """
    page = call_bookstack_api(f"/api/pages/{page_id}")
    html = page.get("html") or ""
    api_attachments = list_attachments_for_page_internal(page_id)
    html_links = extract_attachment_links_from_html(html)
    detected_html_attachment_ids = sorted(
        {
            attachment_id
            for link in html_links
            for attachment_id in link.get("detected_attachment_ids", [])
        }
    )

    return {
        "page_id": page_id,
        "page_name": page.get("name"),
        "clean_page_text_preview": html_to_clean_text(html)[:1000],
        "api_attachment_count": len(api_attachments),
        "api_attachments": api_attachments,
        "html_file_like_links_count": len(html_links),
        "html_file_like_links": html_links,
        "detected_html_attachment_ids": detected_html_attachment_ids,
        "diagnosis": (
            "BookStack API found attachments for this page."
            if api_attachments
            else "No attachments found via /api/attachments for this page. "
                 "If html_file_like_links_count is also 0, the PDF name is only plain text, "
                 "so the server has no file ID/URL to fetch."
        ),
    }


@mcp.tool()
def list_page_attachments(page_id: int) -> dict[str, Any]:
    """
    List all attachments uploaded to a specific BookStack page.
    """
    attachments = list_attachments_for_page_internal(page_id)

    return {
        "page_id": page_id,
        "total_attachments": len(attachments),
        "attachments": attachments,
    }


@mcp.tool()
def get_attachment_text(attachment_id: int) -> dict[str, Any]:
    """
    Test one attachment. Use this first to confirm PDF extraction works.
    """
    return extract_attachment_text(attachment_id)


@mcp.tool()
def list_all_attachments(max_attachments: int = 10000) -> dict[str, Any]:
    """
    Debug tool: shows whether MCP can see BookStack attachments.
    """
    attachments = list_all_attachments_internal(max_attachments=max_attachments)

    return {
        "total_collected": len(attachments),
        "attachments": attachments,
    }


@mcp.tool()
def get_pages_batch_as_rag_text(
    offset: int = 0,
    count: int = 10,
    include_attachments: bool = True,
    max_attachments: int = 1000,
) -> dict[str, Any]:
    """
    Safer Langflow RAG ingestion tool.
    Fetch a small batch of BookStack pages and return clean RAG text.
    Use next_offset repeatedly until has_more=false.
    """
    offset = max(0, offset)
    count = max(1, min(count, 50))
    max_attachments = max(1, min(max_attachments, 10000))

    list_result = call_bookstack_api(
        "/api/pages",
        {
            "count": count,
            "offset": offset,
        },
    )

    page_items = list_result.get("data", [])
    total = list_result.get("total")

    attachments_by_page: dict[int, list[dict[str, Any]]] = {}
    total_attachments_found = 0

    if include_attachments and page_items:
        all_attachments = list_all_attachments_internal(max_attachments=max_attachments)
        total_attachments_found = len(all_attachments)
        attachments_by_page = group_attachments_by_page(all_attachments)

    pages = []
    rag_text_parts = []

    for page_item in page_items:
        page_id = page_item.get("id")
        if page_id is None:
            continue

        page_attachments = attachments_by_page.get(int(page_id), [])
        page = get_page_internal(
            page_id=int(page_id),
            include_attachments=include_attachments,
            page_attachments=page_attachments,
        )
        pages.append(page)
        rag_text_parts.append(format_page_for_rag(page))

    next_offset = offset + len(page_items)
    has_more = bool(isinstance(total, int) and next_offset < total)

    return {
        "offset": offset,
        "count": count,
        "next_offset": next_offset,
        "has_more": has_more,
        "total": total,
        "total_pages_in_batch": len(pages),
        "include_attachments": include_attachments,
        "total_attachments_found": total_attachments_found,
        "total_attachments_added_to_pages": sum(
            page.get("attachment_count", 0) for page in pages
        ),
        "total_attachments_with_text": sum(
            page.get("attachment_text_count", 0) for page in pages
        ),
        "total_attachments_skipped": sum(
            page.get("attachment_skipped_count", 0) for page in pages
        ),
        "pages": [simplify_page(page) for page in pages],
        "rag_text": "\n\n".join(rag_text_parts),
    }


@mcp.tool()
def get_all_pages(
    max_pages: int = 50,
    include_attachments: bool = True,
    max_attachments: int = 1000,
) -> dict[str, Any]:
    """
    Get all BookStack pages.
    If include_attachments=true, uploaded PDF text is added into page text.
    Broken/unreadable attachments are skipped, not allowed to stop the full run.
    """
    max_pages = max(1, min(max_pages, 10000))

    collected_pages = []
    offset = 0
    batch_size = 500

    attachments_by_page: dict[int, list[dict[str, Any]]] = {}
    total_attachments_found = 0

    if include_attachments:
        all_attachments = list_all_attachments_internal(
            max_attachments=max_attachments
        )
        total_attachments_found = len(all_attachments)
        attachments_by_page = group_attachments_by_page(all_attachments)

    while len(collected_pages) < max_pages:
        list_result = call_bookstack_api(
            "/api/pages",
            {
                "count": batch_size,
                "offset": offset,
            },
        )

        page_items = list_result.get("data", [])

        if not page_items:
            break

        for page_item in page_items:
            if len(collected_pages) >= max_pages:
                break

            page_id = page_item.get("id")

            if page_id is None:
                continue

            page_attachments = attachments_by_page.get(int(page_id), [])

            full_page = get_page_internal(
                page_id=int(page_id),
                include_attachments=include_attachments,
                page_attachments=page_attachments,
            )

            collected_pages.append(full_page)

        total = list_result.get("total")

        if isinstance(total, int) and offset + len(page_items) >= total:
            break

        if len(page_items) < batch_size:
            break

        offset += batch_size

    total_attachments_added_to_pages = sum(
        page.get("attachment_count", 0) for page in collected_pages
    )
    total_attachments_with_text = sum(
        page.get("attachment_text_count", 0) for page in collected_pages
    )
    total_attachments_skipped = sum(
        page.get("attachment_skipped_count", 0) for page in collected_pages
    )

    return {
        "total_collected": len(collected_pages),
        "include_attachments": include_attachments,
        "total_attachments_found": total_attachments_found,
        "total_attachments_added_to_pages": total_attachments_added_to_pages,
        "total_attachments_with_text": total_attachments_with_text,
        "total_attachments_skipped": total_attachments_skipped,
        "pages": collected_pages,
    }


@mcp.tool()
def get_all_pages_as_rag_text(
    max_pages: int = 50,
    include_attachments: bool = True,
    max_attachments: int = 1000,
) -> dict[str, Any]:
    """
    Main tool for Langflow RAG.
    Output rag_text should go to Split Text -> Embeddings -> Vector DB.
    """
    result = get_all_pages(
        max_pages=max_pages,
        include_attachments=include_attachments,
        max_attachments=max_attachments,
    )

    pages = result.get("pages", [])
    rag_text_parts = []

    for page in pages:
        rag_text_parts.append(format_page_for_rag(page))

    return {
        "total_pages": len(pages),
        "include_attachments": include_attachments,
        "total_attachments_found": result.get("total_attachments_found"),
        "total_attachments_added_to_pages": result.get("total_attachments_added_to_pages"),
        "total_attachments_with_text": result.get("total_attachments_with_text"),
        "total_attachments_skipped": result.get("total_attachments_skipped"),
        "rag_text": "\n\n".join(rag_text_parts),
    }


async def health_check(request):
    return JSONResponse(
        {
            "status": "ok",
            "service": "bookstack-mcp",
            "bookstack_base_url": BOOKSTACK_BASE_URL,
        }
    )


@contextlib.asynccontextmanager
async def lifespan(app: Starlette):
    async with mcp.session_manager.run():
        yield


app = Starlette(
    routes=[
        Route("/health", health_check),
        Mount("/", app=mcp.streamable_http_app()),
    ],
    lifespan=lifespan,
)


if __name__ == "__main__":
    uvicorn.run(
        app,
        host=MCP_HOST,
        port=MCP_PORT,
    )
